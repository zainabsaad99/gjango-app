from django.shortcuts import render
import json
import os
from docx import Document 
from django.core.mail import send_mail
from dotenv import dotenv_values




# Create your views here.
def read_services(dir):
   # Create an empty list to store service information
   services_folder = []
   
   for d in os.listdir(dir):
      # save name of service folder in a temp dict
      dir_detail_of_services= {"name": d}  
      # Iterate over the dir\service_1 in the firstt iterate 
      for f in os.listdir(os.path.join(dir, d)):
         # f is the list of file name inside each service 
         
         # Check if the file has a ".txt" extension
         if f.endswith(".json"):
            # If it's a text file, open and read its contents
            with open(os.path.join(dir, d, f), 'r', encoding='utf-8') as f:
               # read data from json file according to the language
               txt = json.load(f)
               # store detial in the dictionary 
               dir_detail_of_services["txt"] = txt
         elif f.endswith(".jpg"):
            intro_img = str(os.path.join(dir, d, f))
            # Modify the image path by replacing 'static/' and backslashes with forward slashes
            dir_detail_of_services["intro_img"] = intro_img.replace('static/', '').replace('\\', '/')  
         elif f.endswith(".mp4"):
            intro_img = str(os.path.join(dir, d, f))
            # Modify the image path by replacing 'static/' and backslashes with forward slashes
            dir_detail_of_services["intro_video"] = intro_img.replace('static/', '').replace('\\', '/')     
        


          

    
            
      # Add the temporary dictionary to the list of service data
      services_folder.append(dir_detail_of_services)
   
   # Return the list of service data as the function's result
   return services_folder

def home(request): 
  
   with open('static/text.json', 'r', encoding='utf-8') as f:
      context = json.load(f)
   services = read_services('static/media/services')
   context["services"] = services  
   articles = read_services('static/media/articles')
   context["articles"] = articles 
   achievements = read_services('static/media/achievement')
   context["achievement"] = achievements 
   vid = read_services('static/media/videos')
   context["video"] = vid
 
   return render(request,'main_page/index.html',context)


def service(request,service_name):
   with open('static/text.json', 'r', encoding='utf-8') as f:
      context = json.load(f)

   with open(f'static/media/services/{service_name}/details/description.json', 'r', encoding='utf-8') as f:
               # read data from json file according to the language
              description = json.load(f)
             
   image_list = []
   for image in os.listdir(f'static/media/services/{service_name}/details/media/'):
      if image.endswith('.jpg'):
         image_list.append({"file": f'media/services/{service_name}/details/media/{image}', "type":"img", "title":image.split('.')[0]})
     
   
   context ['service']={"p_name":str(service_name), "desc":description, "image":image_list}   
   return render(request, 'main_page/service.html',context)


def article(request,article_name):
   with open('static/text.json', 'r', encoding='utf-8') as f:
      context = json.load(f)

   with open(f'static/media/articles/{article_name}/details/description.json', 'r', encoding='utf-8') as f:
               # read data from json file according to the language
              description = json.load(f)
             
   image_list = []
   for image in os.listdir(f'static/media/articles/{article_name}/details/media/'):
      if image.endswith('.jpg'):
         image_list.append({"file": f'media/articles/{article_name}/details/media/{image}', "type":"img", "title":image.split('.')[0]})
     
   
  
      # create an instance of a  
      # word document we want to open 
   doc = Document(f'static/media/articles/{article_name}/details/article.docx')
     
   # docdata = [para.text for para in doc.paragraphs]
   docdata = []
   for para in doc.paragraphs:
        para_data = {
            "text": para.text,
            "bold": any(run.bold for run in para.runs),
            "underline": any(run.underline for run in para.runs),
            "font_size": para.style.font.size.pt if para.style.font.size is not None else None,
        }
        docdata.append(para_data)     

   context ['article']={"p_name":str(article_name), "desc":description, "image":image_list,"docdata": docdata}   
   return render(request, 'main_page/article.html',context)

def achievement(request,achievement_name):
   with open('static/text.json', 'r', encoding='utf-8') as f:
      context = json.load(f)

             
   image_list = []
   for image in os.listdir(f'static/media/achievement/{achievement_name}/details/media/'):
      if image.endswith('.jpeg'):
         image_list.append({"file": f'media/achievement/{achievement_name}/details/media/{image}', "type":"img", "title":image.split('.')[0]})
     
   doc = Document(f'static/media/achievement/{achievement_name}/details/achievement.docx')
     
   # docdata = [para.text for para in doc.paragraphs]
   docdata = []
   for para in doc.paragraphs:
        para_data = {
            "text": para.text,
            "bold": any(run.bold for run in para.runs),
            "underline": any(run.underline for run in para.runs),
            "font_size": para.style.font.size.pt if para.style.font.size is not None else None,
        }
        docdata.append(para_data)   
   context ['achievement']={"p_name":str(achievement_name), "docdata": docdata, "image":image_list}   
   print(context)
   return render(request, 'main_page/achievement.html',context)

def contact(request):
    with open('static/text.json', 'r', encoding='utf-8') as f:
      context = json.load(f)
    if request.method == "POST":  
      name = request.POST.get('name')
      email = request.POST.get('email')
      subject = request.POST.get('subject')
      message = request.POST.get('message')
      formatted_message = f"Name: {name}\nEmail: {email}\nSubject: {subject}\n\nMessage:\n{message}"

# Send the email with the formatted message
      send_mail(
         subject,
         formatted_message,
         os.environ.get('EMAIL', ""),  # Use the user's email as "From"
         [os.environ.get('EMAIL_SECOND', "")],  # Use the user-specified recipient email
         fail_silently=False,
         )  
    return render(request, 'main_page/contact.html',context)

def service_detials(request):
    with open('static/text.json', 'r', encoding='utf-8') as f:
      context = json.load(f)
    services = read_services('static/media/services')
    context["services"] = services 
      
    return render(request, 'main_page/service_detail.html',context)

def articles_list(request):
    with open('static/text.json', 'r', encoding='utf-8') as f:
      context = json.load(f)
    articles = read_services('static/media/articles')
    context["articles"] = articles
      
    return render(request, 'main_page/articles_list.html',context)

def achievement_list(request):
    with open('static/text.json', 'r', encoding='utf-8') as f:
      context = json.load(f)
    achievements = read_services('static/media/achievement')
    context["achievement"] = achievements 
      
    return render(request, 'main_page/achievement_list.html',context)


def about(request):
    with open('static/text.json', 'r', encoding='utf-8') as f:
      context = json.load(f)
    return render(request, 'main_page/about.html',context)




def lawyer_info(request):
    with open('static/text.json', 'r', encoding='utf-8') as f:
      context = json.load(f)
      lawyer_data = read_services('static/media/lawyer_info')
      context["lawyerdata"] = lawyer_data  
   
    
    return render(request, 'main_page/lawyer_info.html',context)


def videos(request):
    with open('static/text.json', 'r', encoding='utf-8') as f:
      context = json.load(f)
    vid = read_services('static/media/videos')
    context["video"] = vid

    
    return render(request, 'main_page/videos.html',context)



def lawyer_cv(request,lawyer_name):
   with open('static/text.json', 'r', encoding='utf-8') as f:
      context = json.load(f)

   print(lawyer_name)
   with open(f'static/media/lawyer_info/{lawyer_name}/details/description.json', 'r', encoding='utf-8') as f:
               # read data from json file according to the language
              description = json.load(f)             
 
     
   context ['lawyer']={"p_name":str(lawyer_name), "desc":description} 
   print(context ['lawyer'])
   return render(request, 'main_page/lawyer_cv.html',context)